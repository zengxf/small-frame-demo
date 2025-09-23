import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def generate_addition(carry):
    """生成加法题目，确保结果不超过100"""
    while True:
        a = random.randint(0, 99)
        b = random.randint(0, 99)
        # 确保结果不超过100
        if a + b > 100:
            continue
        if carry:
            if (a % 10 + b % 10) >= 10:
                return a, b
        else:
            if (a % 10 + b % 10) < 10:
                return a, b


def generate_subtraction(borrow):
    """生成减法题目，确保结果为正数"""
    while True:
        a = random.randint(0, 99)
        b = random.randint(0, 99)
        # 确保被减数大于等于减数，结果为正数
        if a < b:
            a, b = b, a
        if borrow:
            if (a % 10) < (b % 10):
                return a, b
        else:
            if (a % 10) >= (b % 10):
                return a, b


# 创建Word文档
doc = Document()
doc.add_heading('小学二年级竖式加减法练习题（800道）', 0)

# 添加说明段落
p = doc.add_paragraph()
p.add_run('说明：本练习包含800道100以内的竖式加减法题目，包括进位加法和退位减法。所有加法结果不超过100。')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 设置表格样式 - 每行4个题目
table = doc.add_table(rows=200, cols=4)
table.style = 'Table Grid'

# 设置表格列宽
for col in table.columns:
    col.width = Pt(150)  # 增加列宽以适应竖式格式

# 生成200道有进位加法、200道无进位加法、200道有退位减法、200道无退位减法
problems = []
problems.extend([('+', a, b) for a, b in [generate_addition(True) for _ in range(200)]])
problems.extend([('+', a, b) for a, b in [generate_addition(False) for _ in range(200)]])
problems.extend([('-', a, b) for a, b in [generate_subtraction(True) for _ in range(200)]])
problems.extend([('-', a, b) for a, b in [generate_subtraction(False) for _ in range(200)]])

# 打乱题目顺序
random.shuffle(problems)

# 填充表格
for i in range(200):  # 200行
    for j in range(4):  # 每行4个题目
        idx = i * 4 + j
        if idx < len(problems):
            op, a, b = problems[idx]
            cell = table.cell(i, j)

            # 清空默认段落
            for paragraph in cell.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # 添加横式（居中）
            p_horizontal = cell.add_paragraph()
            p_horizontal.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_horizontal = p_horizontal.add_run(f"{a} {op} {b} = ")
            run_horizontal.font.size = Pt(12)
            run_horizontal.font.name = '宋体'
            p_horizontal.add_run("\n")

            # 添加竖式（右对齐）
            p_vertical = cell.add_paragraph()
            p_vertical.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 计算最大数字位数，确保对齐
            max_digits = max(len(str(a)), len(str(b)))

            # 第一行：第一个数字（右对齐）
            run1 = p_vertical.add_run(f"{a:>{max_digits}d}")
            run1.font.size = Pt(12)
            run1.font.name = 'Courier New'
            p_vertical.add_run("\n")

            # 第二行：运算符和第二个数字（右对齐）
            run2 = p_vertical.add_run(f"{op}{b:>{max_digits}d}")
            run2.font.size = Pt(12)
            run2.font.name = 'Courier New'
            p_vertical.add_run("\n")

            # 第三行：横线（右对齐）
            run3 = p_vertical.add_run("-" * (max_digits + 1))  # 运算符占一位，所以长度+1
            run3.font.size = Pt(12)
            run3.font.name = 'Courier New'
            p_vertical.add_run("\n\n\n")  # 添加空行，为答案留出空间

            # 设置单元格垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 保存文档
doc.save('二年级竖式加减法练习题.docx')
print("已生成800道题目，保存在'二年级竖式加减法练习题.docx'中。")
print("所有加法题目结果均不超过100。")